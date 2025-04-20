from flask import Flask, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
import os
import tempfile
import re
import docx
import PyPDF2
from pdfminer.high_level import extract_text
from bs4 import BeautifulSoup
import traceback
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Suppress pdfminer warnings
logging.getLogger('pdfminer').setLevel(logging.WARNING)

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), 'resume_uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Enable CORS if needed
from flask_cors import CORS
CORS(app)

# Skills database for extraction
SKILLS_DATABASE = [
    'python', 'java', 'javascript', 'c++', 'sql', 'machine learning',
    'data analysis', 'aws', 'docker', 'git', 'flask', 'django',
    'react', 'angular', 'node.js', 'html', 'css', 'mongodb',
    'postgresql', 'linux', 'rest api', 'graphql'
]

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_from_pdf(file_path):
    """Extract text from PDF using multiple methods with fallbacks"""
    try:
        # First try pdfminer.six
        text = extract_text(file_path)
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"PDFMiner extraction failed: {str(e)}")
    
    try:
        # Fallback to PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            return text if text.strip() else None
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed: {str(e)}")
        return None

def extract_from_docx(file_path):
    """Extract text from DOCX files"""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract paragraphs
        full_text.extend([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return "\n".join(full_text) if full_text else None
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        return None

def extract_name(text):
    """Improved name extraction that prioritizes top-most name"""
    try:
        # First check the very first line before any processing
        first_line = text.split('\n')[0].strip()
        if re.match(r'^[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)+$', first_line):
            return first_line
            
        # Then try other patterns (keep your existing patterns)
        # ...
            
        return None
    except:
        return None

def extract_email(text):
    """Extract email address from text"""
    try:
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        return email_match.group(0) if email_match else None
    except:
        return None
def extract_phone(text):
    """Enhanced to handle Lebanese and international formats"""
    try:
        # Add specific pattern for Lebanese numbers
        leb_pattern = r'(\+?961\s?\d{1,2}\s?\d{3}\s?\d{3})'
        match = re.search(leb_pattern, text)
        if match:
            return match.group(0)
            
    except:
        return None

def extract_skills(text):
    """Extract skills by matching against skills database"""
    try:
        text_lower = text.lower()
        found_skills = []
        
        for skill in SKILLS_DATABASE:
            if skill.lower() in text_lower:
                found_skills.append(skill.title())
        
        return list(set(found_skills))  # Remove duplicates
    except:
        return []

def extract_experience(text):
    """Extract experience information"""
    try:
        # Look for years pattern
        exp_match = re.search(r'(\d+)\s*(?:years?|yrs?|\+)', text, re.IGNORECASE)
        if exp_match:
            return f"{exp_match.group(1)} years"
        
        # Look for experience section
        exp_section = re.search(r'experience(.+?)(?:\n\n|$)', text, re.IGNORECASE | re.DOTALL)
        if exp_section:
            return exp_section.group(1).strip()[:200] + "..."  # Truncate
            
        return None
    except:
        return None

def extract_education(text):
    """Extract education information"""
    try:
        education = []
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in ['university', 'college', 'institute', 'school', 'bachelor', 'master', 'phd', 'degree', 'diploma']):
                # Get the line and possibly the next one
                edu_entry = line
                if i+1 < len(lines):
                    edu_entry += " " + lines[i+1]
                education.append(edu_entry)
        
        return education if education else None
    except:
        return None

def parse_resume_text(text):
    """Parse extracted text into structured data"""
    soup = BeautifulSoup(f"<div>{text}</div>", 'html.parser')
    clean_text = soup.get_text(separator=' ', strip=True)
    
    return {
        "name": extract_name(clean_text),
        "email": extract_email(clean_text),
        "phone": extract_phone(clean_text),
        "skills": extract_skills(clean_text),
        "experience": extract_experience(clean_text),
        "education": extract_education(clean_text)
    }

@app.route('/extract-resume-data/<int:application_id>', methods=['GET'])
def get_resume_data(application_id):
    """Endpoint to get parsed resume data by application ID"""
    try:
        # In a real application, you would fetch the resume path from your database
        # For this example, we'll use a placeholder
        resume_path = f"resumes/{application_id}.pdf"  # Replace with your actual logic
        
        if not os.path.exists(resume_path):
            return jsonify({
                "success": False,
                "message": f"Resume with ID {application_id} not found"
            }), 404
        
        # Extract text based on file type
        text = None
        if resume_path.lower().endswith('.pdf'):
            text = extract_from_pdf(resume_path)
        elif resume_path.lower().endswith(('.doc', '.docx')):
            text = extract_from_docx(resume_path)
        
        if not text:
            return jsonify({
                "success": False,
                "message": "Could not extract text from file"
            }), 400
            
        # Parse the extracted text
        result = parse_resume_text(text)
        return jsonify({
            "success": True,
            "data": result
        })
        
    except Exception as e:
        logger.error(f"Error processing resume: {traceback.format_exc()}")
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/parse', methods=['POST'])
def parse_resume():
    """Endpoint to parse uploaded resume file"""
    if 'resume' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    temp_path = None
    try:
        file = request.files['resume']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
            
        if not allowed_file(file.filename):
            return jsonify({"error": "File type not allowed"}), 400
        
        # Save temporarily
        filename = secure_filename(file.filename)
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(temp_path)
        
        # Extract text based on file type
        text = None
        if filename.lower().endswith('.pdf'):
            text = extract_from_pdf(temp_path)
        elif filename.lower().endswith(('.doc', '.docx')):
            text = extract_from_docx(temp_path)
        
        if not text:
            return jsonify({"error": "Could not extract text from file"}), 400
            
        # Parse the extracted text
        result = parse_resume_text(text)
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error parsing resume: {traceback.format_exc()}")
        return jsonify({
            "error": "An error occurred while processing the resume",
            "details": str(e)
        }), 500
        
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except:
                pass

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)